# -*-coding:utf-8 -*-
# @Time    : 2021/7/5 4:12 下午
# @Author  : Yize Wang
# @File    : docx_operation.py
# @Software: PyCharm

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR_INDEX


class DocxFile:
    """
    create docx file
    """

    # maximum size of the figures
    fig_x_max = 14
    fig_cap_size = 10
    heading_sizes = [14, 14, 12]

    def __init__(self, filename):
        self.filename = filename + ".docx"

        self.doc = Document()
        self.styles = self.doc.styles

        self.num_fig = 0
        self.num_heading = [0 for i in range(len(self.heading_sizes))]

        # add styles
        self.fig_caption = self.styles.add_style("fig_caption", WD_STYLE_TYPE.PARAGRAPH)
        self.fig_caption.font.name = "Times New Roman"
        self.fig_caption.font.size = Pt(self.fig_cap_size)
        self.fig_caption.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self.heading_style = [self.styles.add_style("heading_{}".format(i), WD_STYLE_TYPE.PARAGRAPH) for i in range(3)]
        for i in range(3):
            self.heading_style[i].base_style = self.styles["Heading {}".format(i + 1)]
            self.heading_style[i].font.name = "Times New Roman"
            self.heading_style[i].font.size = Pt(self.heading_sizes[i])
            self.heading_style[i].font.color.theme_color = MSO_THEME_COLOR_INDEX.DARK_1
            self.heading_style[i].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        return

    def add_fig(self, filename, size, title):
        # add this figure
        self.num_fig += 1

        # add figure
        parag = self.doc.add_paragraph()
        parag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = parag.add_run("")
        run.add_picture(filename, width=Cm(self.fig_x_max))
        # add title
        self.doc.add_paragraph("Fig. {}. ".format(self.num_fig) + title, style=self.fig_caption)

        return

    def add_heading(self, context, num):
        # add heading
        self.num_heading[num - 1] += 1

        self.doc.add_paragraph("{}. ".format(self.num_heading[num - 1]) + context, style=self.heading_style[num - 1])

    def __del__(self):
        # save the file
        self.doc.save(self.filename)


if __name__ == '__main__':
    docx = DocxFile("./data/test")

    docx.add_fig("./data/012_results/012_003_Blade_1_fx.png", (8, 4), "012_003_Blade_1_fx")
    docx.add_heading("./data/012_results/012_003_Blade_1_fx.png", 1)